# -*- coding: utf-8 -*-
"""제출 서류 마크다운을 Word(.docx)로 바꾼다.

PDF를 되돌리지 않고 마크다운 원본에서 바로 만든다. PDF 변환기는 표와 이미지를 글자
덩어리로 흩어 놓기 때문에, hwp로 다시 옮길 때 표를 처음부터 다시 그려야 한다.

글꼴은 맑은 고딕이다. 저장소의 Pretendard는 상대 PC에 깔려 있지 않을 수 있고, 없는
글꼴을 지정하면 한글이 네모로 깨진다.

    python scripts/md-to-docx.py <입력.md> <출력.docx>
"""

import io
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "맑은 고딕"
MONO = "D2Coding"
INK = RGBColor(0x19, 0x1F, 0x28)
MUTED = RGBColor(0x64, 0x74, 0x8B)
BLUE = RGBColor(0x0F, 0x54, 0xC9)

# 세로로 긴 휴대전화 캡처(1440x3120)를 두 칸으로 놓았을 때 한 쪽에 들어가는 폭.
IMAGE_WIDTH = Cm(4.6)


def set_font(run, name=FONT, size=None, bold=None, color=None):
    """한글 글꼴은 eastAsia까지 지정해야 적용된다(python-docx는 ascii/hAnsi만 건드린다)."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


TOKEN = re.compile(
    r"(!\[[^\]]*\]\([^)]+\)"  # 이미지
    r"|\[[^\]]+\]\([^)]+\)"  # 링크
    r"|\*\*[^*]+\*\*"  # 굵게
    r"|`[^`]+`"  # 인라인 코드
    r"|\(\s{3,}\))"  # 사람이 채울 빈칸
)


def add_inline(paragraph, text, size=10, color=INK, base_dir="."):
    """한 문단 안의 서식을 조각내어 넣는다. 이미지는 그 자리에 앉힌다."""
    for part in TOKEN.split(text):
        if part == "":
            continue

        image = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", part)
        if image:
            path = os.path.join(base_dir, image.group(2))
            if os.path.exists(path):
                paragraph.add_run().add_picture(path, width=IMAGE_WIDTH)
            continue

        link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", part)
        if link:
            run = paragraph.add_run(link.group(1))
            set_font(run, size=size, color=BLUE)
            run.font.underline = True
            continue

        strong = re.fullmatch(r"\*\*([^*]+)\*\*", part)
        if strong:
            set_font(paragraph.add_run(strong.group(1)), size=size, bold=True, color=color)
            continue

        code = re.fullmatch(r"`([^`]+)`", part)
        if code:
            set_font(paragraph.add_run(code.group(1)), name=MONO, size=size - 0.5, color=color)
            continue

        if re.fullmatch(r"\(\s{3,}\)", part):
            # 손으로 채울 자리. 밑줄만 남겨 인쇄해도 쓸 수 있게 한다.
            run = paragraph.add_run("　" * 6)
            set_font(run, size=size)
            run.font.underline = True
            continue

        set_font(paragraph.add_run(part), size=size, color=color)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_divider(line):
    return re.fullmatch(r"\|[\s:\-|]+\|", line.strip()) is not None


def build(doc, lines, base_dir):
    index = 0
    total = len(lines)

    while index < total:
        raw = lines[index]
        line = raw.strip()

        if line == "":
            index += 1
            continue

        # 수평선은 문서 구분자로만 쓰였다. Word에서는 빈 줄로 충분하다.
        if re.fullmatch(r"-{3,}", line):
            index += 1
            continue

        heading = re.match(r"(#{1,6})\s+(.*)", line)
        if heading:
            level = len(heading.group(1))
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(14 if level > 1 else 0)
            paragraph.paragraph_format.space_after = Pt(6)
            sizes = {1: 17, 2: 13, 3: 11.5}
            add_inline(
                paragraph,
                heading.group(2).replace("**", ""),
                size=sizes.get(level, 11),
                base_dir=base_dir,
            )
            for run in paragraph.runs:
                run.font.bold = True
            index += 1
            continue

        # 코드펜스
        if line.startswith("```"):
            index += 1
            block = []
            while index < total and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            index += 1
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5)
            set_font(paragraph.add_run("\n".join(block)), name=MONO, size=9, color=MUTED)
            continue

        # 표
        if line.startswith("|") and index + 1 < total and is_divider(lines[index + 1]):
            header = split_row(line)
            index += 2
            rows = []
            while index < total and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index].strip()))
                index += 1

            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 캡처를 나란히 놓으려고 쓴 `| | |` 처럼 머리글이 비어 있는 표가 있다.
            # 그대로 두면 빈 머리 줄이 한 칸 생겨 표가 어긋나 보인다.
            if any(value != "" for value in header):
                for position, value in enumerate(header):
                    cell = table.rows[0].cells[position]
                    cell.text = ""
                    add_inline(cell.paragraphs[0], value, size=9.5, base_dir=base_dir)
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
            else:
                table._tbl.remove(table.rows[0]._tr)

            for values in rows:
                cells = table.add_row().cells
                for position, value in enumerate(values[: len(header)]):
                    cells[position].text = ""
                    add_inline(
                        cells[position].paragraphs[0], value, size=9.5, base_dir=base_dir
                    )
            doc.add_paragraph()
            continue

        # 인용은 들여쓴 회색 문단으로 옮긴다.
        if line.startswith(">"):
            block = []
            while index < total and lines[index].strip().startswith(">"):
                block.append(re.sub(r"^\s*>\s?", "", lines[index]).strip())
                index += 1
            for chunk in " \n".join(block).split("\n"):
                if chunk.strip() == "":
                    continue
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.4)
                paragraph.paragraph_format.space_after = Pt(2)
                add_inline(paragraph, chunk.strip(), size=9.5, color=MUTED, base_dir=base_dir)
            continue

        # 목록. 항목 사이 빈 줄은 목록을 끊지 않는다.
        if re.match(r"\s*(?:[-*]|\d+\.)\s+", raw):
            ordered = re.match(r"\s*\d+\.\s+", raw) is not None
            number = 0
            while index < total:
                current = lines[index]
                if current.strip() == "":
                    lookahead = index + 1
                    while lookahead < total and lines[lookahead].strip() == "":
                        lookahead += 1
                    if lookahead < total and re.match(
                        r"\s*(?:[-*]|\d+\.)\s+", lines[lookahead]
                    ):
                        index = lookahead
                        continue
                    break
                head = re.match(r"\s*(?:[-*]|\d+\.)\s+(.*)", current)
                if head is None:
                    break
                body = [head.group(1)]
                index += 1
                while (
                    index < total
                    and lines[index].strip() != ""
                    and re.match(r"\s{2,}\S", lines[index])
                    and not re.match(r"\s*(?:[-*]|\d+\.)\s+", lines[index])
                ):
                    body.append(lines[index].strip())
                    index += 1

                text = " ".join(body)
                number += 1
                checkbox = re.match(r"\[([ xX])\]\s*(.*)", text)
                if checkbox:
                    marker = "☑ " if checkbox.group(1).lower() == "x" else "☐ "
                    text = checkbox.group(2)
                elif ordered:
                    marker = "%d. " % number
                else:
                    marker = "· "

                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.6)
                paragraph.paragraph_format.space_after = Pt(2)
                set_font(paragraph.add_run(marker), size=10)
                add_inline(paragraph, text, size=10, base_dir=base_dir)
            continue

        # 문단
        block = []
        while index < total and lines[index].strip() != "":
            nxt = lines[index].strip()
            if nxt.startswith(("#", ">", "|", "```")) or re.fullmatch(r"-{3,}", nxt):
                break
            if re.match(r"\s*(?:[-*]|\d+\.)\s+", lines[index]):
                break
            block.append(nxt)
            index += 1
        if block:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline(paragraph, " ".join(block), base_dir=base_dir)
        else:
            index += 1


def main():
    source, target = sys.argv[1], sys.argv[2]
    base_dir = os.path.dirname(os.path.abspath(source))
    text = io.open(source, encoding="utf-8").read()

    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    for side in ("left_margin", "right_margin"):
        setattr(section, side, Cm(1.8))
    section.top_margin = section.bottom_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    build(doc, text.split("\n"), base_dir)
    doc.save(target)
    print("%s -> %s (%d KB)" % (source, target, os.path.getsize(target) // 1024))


if __name__ == "__main__":
    main()
